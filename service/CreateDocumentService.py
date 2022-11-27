from docx import Document
from docx.shared import Pt

from service import Constant


def setFont(run, **params):
    f = run.font
    for key, value in params.items():
        if key == "name":
            f.name = value
        if key == "size":
            f.size = Pt(value)
        if key == "color":
            f.color.rgb = value
        if key == "bold":
            f.bold = value
        if key == "italic":
            f.italic = value


def createParagraphHeader(document, text, style):
    pg1 = document.add_paragraph()
    run1 = pg1.add_run(text)
    setFont(run1, size=12, bold=True, name=Constant.FONT_ARIAL)
    pg1.style = style


def createParagraphBody(document, text, style, isItalic, fontName):
    pg1 = document.add_paragraph()
    run1 = pg1.add_run(text)
    setFont(run1, size=12, italic=isItalic, name=fontName)
    pg1.style = style


def createDocument(dataDocument):
    document = Document()

    document.add_heading("Deploy \"" + dataDocument.projectName + "\"", 0)

    p = document.add_paragraph()
    run = p.add_run('Progetto GCloud: ')
    setFont(run, size=14, color=Constant.HEADER_COLOR, bold=True)
    t = p.add_run(Constant.GOOGLE_CLOUD_NAME)
    setFont(t, size=12, color=Constant.BLACK_COLOR, bold=True, name=Constant.FONT_ARIAL)

    p1 = document.add_paragraph()
    run = p1.add_run('Cluster:')
    setFont(run, size=14, color=Constant.HEADER_COLOR, bold=True, name=Constant.FONT_ARIAL)
    t1 = p1.add_run(Constant.CLUSTER_NAME)
    setFont(t1, size=12, color=Constant.BLACK_COLOR, bold=True, name=Constant.FONT_ARIAL)

    document.add_paragraph()

    run = document.add_paragraph().add_run('Istruzioni per il deploy in produzione\n')
    setFont(run, size=14, name=Constant.FONT_ARIAL)

    # P1
    createParagraphHeader(document, 'Docker pull dell’immagine dall’ambiente di test:', 'List Number 2')

    createParagraphBody(document,
                        'eu.gcr.io/tlp-public-api-uservices-test/tlp-digital-java/' + dataDocument.serviceName + ':' + dataDocument.serviceVersion,
                        'List Continue 2', True, Constant.FONT_TIMES_NEW_ROMAN)
    createParagraphBody(document,
                        'Effettuare il tag e push su registry di produzione:', 'List Continue 2', False,
                        Constant.FONT_ARIAL)
    createParagraphBody(document,
                        'eu.gcr.io/tlp-public-api-uservices-prod/tlp-digital-java/' + dataDocument.serviceName + ':' + dataDocument.serviceVersion,
                        'List Continue 2', True, Constant.FONT_TIMES_NEW_ROMAN)
    document.add_paragraph()

    # P2
    createParagraphHeader(document, 'Recuperare i file di deploy da:', 'List Number 2')

    createParagraphBody(document,
                        'https://source.cloud.google.com/tlp-shared-services/tlp-java-uservices-gke-deploy/+/master:' + dataDocument.gkeServiceName + '/',
                        'List Continue 2', False, Constant.FONT_TIMES_NEW_ROMAN)
    document.add_paragraph()

    # P3
    createParagraphHeader(document, 'Eseguire i seguenti comandi rispettando l’ordine:', 'List Number 2')

    for command in dataDocument.commandList:
        createParagraphBody(document, command, 'List Continue 2', True, Constant.FONT_TIMES_NEW_ROMAN)
    document.add_paragraph()

    document.save(Constant.PARENT_DIR + "/" + dataDocument.directoryName + "/" + dataDocument.fileName)
